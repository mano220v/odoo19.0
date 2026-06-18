# -*- coding: utf-8 -*-
import io
import json
import logging
import re
from html import escape

from odoo import http, _
from odoo.http import request, Response

_logger = logging.getLogger(__name__)


def _strip_html(html_content):
    """Strip HTML tags and return plain text."""
    text = re.sub(r'<[^>]+>', ' ', html_content or '')
    return re.sub(r'\s+', ' ', text).strip()


class WordEditorController(http.Controller):

    # ─────────────────────────────────────────────────────────────────────────
    # DOCX Export
    # ─────────────────────────────────────────────────────────────────────────

    @http.route('/word_editor/export/docx/<int:doc_id>', type='http', auth='user', methods=['GET'])
    def export_docx(self, doc_id, **kwargs):
        document = request.env['word.document'].browse(doc_id)
        if not document.exists():
            return request.not_found()
        document.check_access_rights('read')
        document.check_access_rule('read')

        try:
            from docx import Document as DocxDocument
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from bs4 import BeautifulSoup, NavigableString, Tag

            doc = DocxDocument()

            # ── Page Setup (A4) ──────────────────────────────────────────────
            section = doc.sections[0]
            section.page_width  = Inches(8.27)   # A4 width
            section.page_height = Inches(11.69)  # A4 height
            section.left_margin   = Inches(1.0)
            section.right_margin  = Inches(1.0)
            section.top_margin    = Inches(1.0)
            section.bottom_margin = Inches(1.0)

            # ── Document Properties ──────────────────────────────────────────
            props = doc.core_properties
            props.title  = document.name
            props.author = document.author_id.name or ''

            # ── HTML → DOCX Conversion ───────────────────────────────────────
            soup = BeautifulSoup(document.content or '<p></p>', 'html.parser')

            def _apply_run_style(run, element):
                """Apply inline styles from an HTML element to a docx Run."""
                style_attr = element.get('style', '')
                parent_tags = {p.name for p in element.parents}
                if element.name in ('strong', 'b') or 'strong' in parent_tags or 'b' in parent_tags:
                    run.bold = True
                if element.name in ('em', 'i') or 'em' in parent_tags or 'i' in parent_tags:
                    run.italic = True
                if element.name in ('u',) or 'u' in parent_tags:
                    run.underline = True
                if element.name in ('s', 'strike', 'del') or any(t in parent_tags for t in ('s', 'strike', 'del')):
                    run.font.strike = True
                if 'font-weight: bold' in style_attr or 'font-weight:bold' in style_attr:
                    run.bold = True
                if 'font-style: italic' in style_attr or 'font-style:italic' in style_attr:
                    run.italic = True
                # Font size
                size_match = re.search(r'font-size:\s*(\d+)pt', style_attr)
                if size_match:
                    run.font.size = Pt(int(size_match.group(1)))

            def _add_paragraph_from_tag(tag, doc_obj, para=None):
                """Recursively build a docx paragraph from an HTML tag."""
                if para is None:
                    para = doc_obj.add_paragraph()

                for child in tag.children:
                    if isinstance(child, NavigableString):
                        text = str(child)
                        if text.strip():
                            run = para.add_run(text)
                            _apply_run_style(run, tag)
                    elif isinstance(child, Tag):
                        if child.name in ('br',):
                            run = para.add_run('\n')
                        elif child.name in ('strong', 'b', 'em', 'i', 'u', 's', 'span', 'a', 'mark'):
                            for grandchild in child.children:
                                if isinstance(grandchild, NavigableString):
                                    text = str(grandchild)
                                    if text:
                                        run = para.add_run(text)
                                        _apply_run_style(run, child)
                                elif isinstance(grandchild, Tag):
                                    _add_paragraph_from_tag(grandchild, doc_obj, para)
                        else:
                            _add_paragraph_from_tag(child, doc_obj, para)
                return para

            heading_map = {
                'h1': ('Heading 1', 1),
                'h2': ('Heading 2', 2),
                'h3': ('Heading 3', 3),
                'h4': ('Heading 4', 4),
                'h5': ('Heading 5', 5),
                'h6': ('Heading 6', 6),
            }

            title_para = doc.add_heading(document.name, level=0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()  # spacer

            def _process_element(el, doc_obj):
                tag_name = el.name
                if tag_name in heading_map:
                    doc_obj.add_heading(el.get_text(strip=True), level=heading_map[tag_name][1])
                elif tag_name == 'p':
                    _add_paragraph_from_tag(el, doc_obj)
                elif tag_name in ('ul', 'ol'):
                    style = 'List Bullet' if tag_name == 'ul' else 'List Number'
                    for li in el.find_all('li', recursive=False):
                        p = doc_obj.add_paragraph(style=style)
                        for child in li.children:
                            if isinstance(child, NavigableString):
                                p.add_run(str(child))
                            elif isinstance(child, Tag):
                                _add_paragraph_from_tag(child, doc_obj, p)
                elif tag_name == 'table':
                    rows_data = el.find_all('tr')
                    if rows_data:
                        cols = max(len(r.find_all(['td', 'th'])) for r in rows_data)
                        if cols:
                            tbl = doc_obj.add_table(rows=len(rows_data), cols=cols)
                            tbl.style = 'Table Grid'
                            for ri, row in enumerate(rows_data):
                                cells = row.find_all(['td', 'th'])
                                for ci, cell in enumerate(cells):
                                    if ci < cols:
                                        tbl.cell(ri, ci).text = cell.get_text(strip=True)
                elif tag_name in ('div', 'section', 'article', 'main', 'body'):
                    for child in el.children:
                        if isinstance(child, Tag):
                            _process_element(child, doc_obj)
                        elif isinstance(child, NavigableString) and child.strip():
                            doc_obj.add_paragraph(str(child).strip())
                elif tag_name == 'hr':
                    doc_obj.add_paragraph('-' * 60)
                elif tag_name == 'blockquote':
                    p = doc_obj.add_paragraph(el.get_text(strip=True))
                    p.style.font.italic = True

            for child in soup.children:
                if isinstance(child, Tag):
                    _process_element(child, doc)
                elif isinstance(child, NavigableString) and child.strip():
                    doc.add_paragraph(str(child).strip())

            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)

            filename = re.sub(r'[^\w\-_\. ]', '_', document.name) + '.docx'
            content_type = (
                'application/vnd.openxmlformats-officedocument.'
                'wordprocessingml.document'
            )
            return Response(
                buf.getvalue(),
                headers={
                    'Content-Type': content_type,
                    'Content-Disposition': f'attachment; filename="{filename}"',
                    'Content-Length': str(buf.getbuffer().nbytes),
                }
            )

        except ImportError as e:
            _logger.warning('python-docx / beautifulsoup4 not installed: %s', e)
            # Fallback: export as HTML with .docx-like content-type
            html_content = f"""<!DOCTYPE html>
<html>
<head>
  <meta charset="UTF-8">
  <title>{escape(document.name)}</title>
  <style>
    body {{ font-family: Arial, sans-serif; margin: 2.54cm; font-size: 12pt; }}
    h1,h2,h3 {{ color: #1a1a1a; }}
    table {{ border-collapse: collapse; }}
    td, th {{ border: 1px solid #000; padding: 6px 12px; }}
  </style>
</head>
<body>
  <h1>{escape(document.name)}</h1>
  {document.content or ''}
</body>
</html>"""
            filename = re.sub(r'[^\w\-_\. ]', '_', document.name) + '.html'
            return Response(
                html_content.encode('utf-8'),
                headers={
                    'Content-Type': 'text/html; charset=utf-8',
                    'Content-Disposition': f'attachment; filename="{filename}"',
                    'X-Export-Note': (
                        'Install python-docx and beautifulsoup4 for native .docx export'
                    ),
                }
            )

        except Exception as e:
            _logger.error('DOCX export failed for doc %s: %s', doc_id, e)
            return Response(
                f'Export failed: {str(e)}'.encode(),
                status=500,
                headers={'Content-Type': 'text/plain'},
            )

    # ─────────────────────────────────────────────────────────────────────────
    # PDF Export
    # ─────────────────────────────────────────────────────────────────────────

    @http.route('/word_editor/export/pdf/<int:doc_id>', type='http', auth='user', methods=['GET'])
    def export_pdf(self, doc_id, **kwargs):
        document = request.env['word.document'].browse(doc_id)
        if not document.exists():
            return request.not_found()
        document.check_access_rights('read')
        document.check_access_rule('read')
        try:
            pdf_content, _ = request.env['ir.actions.report']._render_qweb_pdf(
                'word_editor.report_word_document_pdf', [doc_id]
            )
            filename = re.sub(r'[^\w\-_\. ]', '_', document.name) + '.pdf'
            return Response(
                pdf_content,
                headers={
                    'Content-Type': 'application/pdf',
                    'Content-Disposition': f'attachment; filename="{filename}"',
                    'Content-Length': str(len(pdf_content)),
                }
            )
        except Exception as e:
            _logger.error('PDF export failed for doc %s: %s', doc_id, e)
            return Response(
                f'PDF export failed: {str(e)}'.encode(),
                status=500,
                headers={'Content-Type': 'text/plain'},
            )

    # ─────────────────────────────────────────────────────────────────────────
    # Auto-Save API (JSON-RPC)
    # ─────────────────────────────────────────────────────────────────────────

    @http.route('/word_editor/autosave', type='jsonrpc', auth='user', methods=['POST'])
    def autosave(self, doc_id=None, name=None, content=None, **kwargs):
        try:
            vals = {}
            if name is not None:
                vals['name'] = name
            if content is not None:
                vals['content'] = content

            if doc_id:
                doc = request.env['word.document'].browse(int(doc_id))
                if doc.exists():
                    doc.check_access_rights('write')
                    doc.check_access_rule('write')
                    doc.write(vals)
                    return {'status': 'ok', 'doc_id': doc.id}
            else:
                if not vals.get('name'):
                    vals['name'] = 'Untitled Document'
                doc = request.env['word.document'].create(vals)
                return {'status': 'created', 'doc_id': doc.id}
        except Exception as e:
            _logger.error('Auto-save failed: %s', e)
            return {'status': 'error', 'message': str(e)}

    # ─────────────────────────────────────────────────────────────────────────
    # Template List API
    # ─────────────────────────────────────────────────────────────────────────

    @http.route('/word_editor/templates', type='jsonrpc', auth='user', methods=['POST'])
    def get_templates(self, **kwargs):
        templates = request.env['word.document.template'].search_read(
            [('active', '=', True)],
            ['name', 'category', 'description', 'content', 'icon', 'color'],
            order='sequence, name',
        )
        return templates

    # ─────────────────────────────────────────────────────────────────────────
    # Public share view (read-only)
    # ─────────────────────────────────────────────────────────────────────────

    @http.route('/word_editor/view/<int:doc_id>', type='http', auth='public', website=False)
    def view_shared(self, doc_id, **kwargs):
        document = request.env['word.document'].sudo().browse(doc_id)
        if not document.exists() or not document.is_shared:
            return request.not_found()

        html = f"""<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0">
  <title>{escape(document.name)}</title>
  <style>
    * {{ box-sizing: border-box; margin: 0; padding: 0; }}
    body {{ background: #f5f5f5; font-family: Arial, sans-serif; padding: 40px 20px; }}
    .page {{ background: #fff; max-width: 800px; margin: 0 auto;
             padding: 2.54cm; box-shadow: 0 2px 16px rgba(0,0,0,.15); }}
    h1.doc-title {{ font-size: 1.6em; color: #333; margin-bottom: 24px;
                    border-bottom: 2px solid #4a90d9; padding-bottom: 8px; }}
    .meta {{ color: #666; font-size: 0.85em; margin-bottom: 24px; }}
    .content {{ line-height: 1.6; color: #222; }}
    .content img {{ max-width: 100%; }}
    .content table {{ border-collapse: collapse; width: 100%; }}
    .content td, .content th {{ border: 1px solid #ddd; padding: 8px; }}
  </style>
</head>
<body>
  <div class="page">
    <h1 class="doc-title">{escape(document.name)}</h1>
    <div class="meta">
      By {escape(document.author_id.name or '')} &bull;
      {document.write_date.strftime('%B %d, %Y') if document.write_date else ''}
      &bull; {document.word_count} words
    </div>
    <div class="content">{document.content or ''}</div>
  </div>
</body>
</html>"""
        return Response(html, headers={'Content-Type': 'text/html; charset=utf-8'})
